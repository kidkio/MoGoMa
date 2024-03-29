import pytesseract
import random
import docx
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import smtplib
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import os
import cv2 as cv
import cv2
import numpy as np
from flask import Flask,request,render_template,send_file,send_from_directory,Response
from PIL import Image
from werkzeug.utils import secure_filename
import os
import requests

def kor2eng(query):
    url = "https://translate.kakao.com/translator/translate.json"

    headers = {
        "Referer": "https://translate.kakao.com/",
        "User-Agent": "Mozilla/5.0"
    }

    data = {
        "queryLanguage": "kr",
        "resultLanguage": "en",
        "q": query
    }

    resp = requests.post(url, headers=headers, data=data)
    data = resp.json()
    output = data['result']['output'][0][0]
    return output
    
def eng2kor(query):
    url = "https://translate.kakao.com/translator/translate.json"

    headers = {
        "Referer": "https://translate.kakao.com/",
        "User-Agent": "Mozilla/5.0"
    }

    data = {
        "queryLanguage": "en",
        "resultLanguage": "kr",
        "q": query
    }

    resp = requests.post(url, headers=headers, data=data)
    data = resp.json()
    output = data['result']['output'][0][0]
    return output



app = Flask("MoGoMa",static_url_path='')
app.config['UPLOAD_FOLDER'] = "/mnt/c/Users/am030/Documents/MoGoMa/uploads/"
app.config['RESULT_FORDER'] = "/mnt/c/Users/am030/Documents/MoGoMa/result/"

def divide_list(l, n): 
    # 리스트 l의 길이가 n이면 계속 반복
    for i in range(0, len(l), n): 
        yield l[i:i + n]        

        
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/help")
def help():
    return render_template("help.html")

@app.route("/text_mode/download",methods=['POST','GET'])
def tm_download():
    if request.method == "POST":
        text = request.form['text']
        file_name = request.form['file_name']
        print(text)
    text_processing(text=text,title=file_name)
    return send_from_directory(app.config['RESULT_FORDER'],file_name+".docx",as_attachment=True)



right = False
@app.route("/download",methods=['POST','GET'])
def download():
    global right
    if request.method == "POST":
        file = request.files['file']
        file_path = os.path.join(app.config['UPLOAD_FOLDER'],secure_filename(file.filename))
        file.save(file_path)
    s = os.path.splitext(file_path)
    s = os.path.split(s[0])
    right = False
    print(s[1])
    error_handler(title=s[1],file_path=file_path)
    print(right)
    if right:
        return send_from_directory(app.config['RESULT_FORDER'],s[1]+".docx", as_attachment=True)
    else:
        return render_template("unbound.html")
    
@app.route("/text_mode")
def text_mode():
    return render_template("text_mode.html")


def error_handler(title,file_path):
    global right
    try:
        processing(title=title,file_path=file_path)
        right = True
        print("right in error_errorhandler :",right)
    except AttributeError:
        print("nonetype error")
        error_handler(title=title,file_path=file_path)
    except UnboundLocalError:
        right = False

def text_processing(text,title):
    sent = text
    Fsent = []
    K_sents = []
    Random_E_sents = []
    E_sents = []
    Answer_Word_Finish = []
    Word_Finish = [[]]
    NoticeAble = []
    Fsent = sent.split(".")
    del Fsent[-1]
    for text in Fsent:
        text = text.replace("\n",' ')
        E_sents.append(text)
        Random_E_sents.append(text)
        print(text)
#         result = translator.translate(text, src="en", dest="ko")
#         result = translate_client.translate(text, target_language="ko")
        result = eng2kor(text)
        #print(result)
        K_sents.append(result)
        #sents = str(sents).replace(",","")
        sents = text.split(" ")
        Word_Finish.append(sents)
    for textt in Fsent:
        textt = textt.replace("\n",' ')
        sent = textt.split(" ")
        Answer_Word_Finish.append(sent)
    del Word_Finish[0]
    #단어 재배열
    for i in range(0,int(len(K_sents))):
        #print(K_sents[i])
        random.shuffle(Word_Finish[i])
        #print(Word_Finish[i])
    print(Answer_Word_Finish, "단어 답지\n")
    print(Word_Finish, "단어 배열\n")
    print(K_sents, "문장 해석\n")
    print(E_sents, "영어 문장 답지\n")
    results = list(divide_list(E_sents, 3))
    random.shuffle(results)
    random.shuffle(Random_E_sents)
    print(Random_E_sents, "랜덤 문장 문제\n")
    Finish = []
    Answer_Finish = []
    for i in range(1, len(K_sents)):
        Finish.append(str(i) + "." + K_sents[i-1])
        Finish.append(Word_Finish[i-1])
    for i in range(1, len(K_sents)):
        Answer_Finish.append(str(i) + "." + K_sents[i-1])
        Answer_Finish.append(Answer_Word_Finish[i-1])
    for words in Word_Finish:                                                 # 어려운 단어 추가??
        #print("================================================================")
        for word_2 in words:
            if len(str(word_2)) > 10:
                if ',' in str(word_2) or '\'' in str(word_2):
                    pass
                else:
                    try:
                        print("===========================================")
                        NoticeKor = eng2kor(word_2)
                        #print(NoticeKor)
                        if NoticeKor.encode().isalpha():
                            pass
                        else:
                            print(NoticeKor)
                            print(word_2)
                            #word_2 = word_2.replace(' ','')
                            if word_2 not in NoticeAble:
                                NoticeAble.append(word_2)
                                NoticeAble.append(NoticeKor)
                    except:
                        pass

    #for i in range(1, len(Word_Finish)):
    #    if len(Word_Finish[i] > 6):
    #        NoticeAble.append(Word_Finish[i])
    #        NoticeAble.append(translator.translate(Word_Finish[i], src="en", dest="ko"))



    doc = docx.Document()
    paragraph = doc.add_paragraph('-우석쌤 잘가요-')
    doc.add_heading ( ' 영어 단어 랜덤 배열 ' , 0)
    para = doc.add_paragraph()
    for number in range(0, len(Finish)):
        Finish[number] = str(Finish[number]).replace('\'','')
        Finish[number] = str(Finish[number]).replace('\"','')
        run = para.add_run(str(Finish[number]))
        para.add_run("\n")
        if number % 2 == 1:
            para.add_run("\n")

    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 단어 랜덤 배열 답지' , 0)
    para = doc.add_paragraph()
    for number in range(0,len(Answer_Finish)):
        Answer_Finish[number] = str(Answer_Finish[number]).replace('\'','')
        Answer_Finish[number] = str(Answer_Finish[number]).replace('\"','')
        run = para.add_run(str(Answer_Finish[number]))
        para.add_run("\n")
        #if number % 2 == 1:
            #para.add_run("\n")
    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열' , 0)
    para = doc.add_paragraph()
    for number in range(0,len(Random_E_sents)):
        run = para.add_run(str(Random_E_sents[number]))
        para.add_run("\n\n")
    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열-3' , 0)
    para = doc.add_paragraph()
    for number in range(0,len(results)):
        #results = str(results[number]).replace('\'','')
        run = para.add_run(str(results[number]))
        para.add_run("\n\n\n")
    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열-답지' , 0)
    para = doc.add_paragraph()
    for number in range(0,len(E_sents)):
        run = para.add_run(str(E_sents[number]))
        para.add_run("\n\n")
    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 어려운 영어 단어 모음집' , 0)
    para = doc.add_paragraph()
    for number in range(0,len(NoticeAble)):
        run = para.add_run(str(NoticeAble[number]))
        para.add_run("\n")
    run.font.name = 'Segoe UI Light'
    run.bold = True
    run.font.size = docx.shared.Pt(9)
    doc.save("/mnt/c/Users/am030/Documents/MoGoMa/result/"+f"{title}.docx")
    print("processing finished!")

    
def processing(title,file_path):
    sent = pytesseract.image_to_string(Image.open(file_path))

#     translator = Translator()

    Fsent = []
    K_sents = []
    Random_E_sents = []
    E_sents = []
    Answer_Word_Finish = []
    Word_Finish = [[]]
    NoticeAble = []
    Fsent = sent.split(".")
    del Fsent[-1]
    for text in Fsent:
        text = text.replace("\n",' ')
        E_sents.append(text)
        Random_E_sents.append(text)
        print(text)
#         result = translator.translate(text, src="en", dest="ko")
#         result = translate_client.translate(text, target_language="ko")
        result = eng2kor(text)
        #print(result)
        K_sents.append(result)
        #sents = str(sents).replace(",","")
        sents = text.split(" ")
        Word_Finish.append(sents)

    for textt in Fsent:
        textt = textt.replace("\n",' ')
        sent = textt.split(" ")

        Answer_Word_Finish.append(sent)


    del Word_Finish[0]
    #단어 재배열

    for i in range(0,int(len(K_sents))):
        #print(K_sents[i])
        random.shuffle(Word_Finish[i])
        #print(Word_Finish[i])

    print(Answer_Word_Finish, "단어 답지\n")    
    print(Word_Finish, "단어 배열\n")
    print(K_sents, "문장 해석\n")
    print(E_sents, "영어 문장 답지\n")

    results = list(divide_list(E_sents, 3))
    random.shuffle(results) 

    random.shuffle(Random_E_sents)
    print(Random_E_sents, "랜덤 문장 문제\n")

    Finish = []
    Answer_Finish = []
    for i in range(1, len(K_sents)):
        Finish.append(str(i) + "." + K_sents[i-1])
        Finish.append(Word_Finish[i-1])

    for i in range(1, len(K_sents)):
        Answer_Finish.append(str(i) + "." + K_sents[i-1])
        Answer_Finish.append(Answer_Word_Finish[i-1])


    for words in Word_Finish:                                                 # 어려운 단어 추가??
        #print("================================================================")
        for word_2 in words:
            if len(str(word_2)) > 10:
                if ',' in str(word_2) or '\'' in str(word_2):
                    pass
                else:
                    try:
                        print("===========================================")
                        NoticeKor = eng2kor(word_2)
                        #print(NoticeKor)
                        if NoticeKor.encode().isalpha():
                            pass
                        else:
                            print(NoticeKor)
                            print(word_2)
                            #word_2 = word_2.replace(' ','')
                            if word_2 not in NoticeAble:
                                NoticeAble.append(word_2)
                                NoticeAble.append(NoticeKor)
                    except:
                        pass

    #for i in range(1, len(Word_Finish)):
    #    if len(Word_Finish[i] > 6):
    #        NoticeAble.append(Word_Finish[i])
    #        NoticeAble.append(translator.translate(Word_Finish[i], src="en", dest="ko"))



    doc = docx.Document()
    paragraph = doc.add_paragraph('-우석쌤 잘가요-')
    doc.add_heading ( ' 영어 단어 랜덤 배열 ' , 0)
    para = doc.add_paragraph()
    for number in range(0, len(Finish)):
        Finish[number] = str(Finish[number]).replace('\'','')
        Finish[number] = str(Finish[number]).replace('\"','')
        run = para.add_run(str(Finish[number]))
        para.add_run("\n")
        if number % 2 == 1:
            para.add_run("\n")



    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 단어 랜덤 배열 답지' , 0)                       #가독성을 위해서 E_sents로 바꿀 수 있음 <피드백필요>
    para = doc.add_paragraph()
    for number in range(0,len(Answer_Finish)):
        Answer_Finish[number] = str(Answer_Finish[number]).replace('\'','')
        Answer_Finish[number] = str(Answer_Finish[number]).replace('\"','')
        run = para.add_run(str(Answer_Finish[number]))
        para.add_run("\n")
        #if number % 2 == 1:
            #para.add_run("\n")

    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열' , 0)                       #문장을 몇개씩 묶어서 기능할 수 있음 <피드백필요>
    para = doc.add_paragraph()
    for number in range(0,len(Random_E_sents)):
        run = para.add_run(str(Random_E_sents[number]))
        para.add_run("\n\n")

    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열-3' , 0)                       #문장을 몇개씩 묶어서 기능할 수 있음 <피드백필요>
    para = doc.add_paragraph()
    for number in range(0,len(results)):
        #results = str(results[number]).replace('\'','')
        run = para.add_run(str(results[number]))
        para.add_run("\n\n\n")

    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 영어 문장 랜덤 배열-답지' , 0)                       #문장을 몇개씩 묶어서 기능할 수 있음 <피드백필요>
    para = doc.add_paragraph()
    for number in range(0,len(E_sents)):
        run = para.add_run(str(E_sents[number]))
        para.add_run("\n\n")

    run.add_break(WD_BREAK.PAGE)
    doc.add_heading ( ' 어려운 영어 단어 모음집' , 0)                       #문장을 몇개씩 묶어서 기능할 수 있음 <피드백필요>
    para = doc.add_paragraph()
    for number in range(0,len(NoticeAble)):
        run = para.add_run(str(NoticeAble[number]))
        para.add_run("\n")

    run.font.name = 'Segoe UI Light'
    run.bold = True
    run.font.size = docx.shared.Pt(9)
    doc.save("/mnt/c/Users/am030/Documents/MoGoMa/result/"+f"{title}.docx")
    print("processing finished!")

app.run(host="0.0.0.0",port=5000,debug=True)
